import csv
from collections.abc import Callable
from datetime import date, datetime, timezone
from io import BytesIO, StringIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from sqlmodel import Session

from app.repositories import reports as report_repo
from app.schemas.report import DisciplineRow, ReportTaskRow

EXPORT_TEMPLATE_FULL = "full"
EXPORT_TEMPLATE_COMPACT = "compact"
REPORT_COLUMNS_FULL = [
    "task_id",
    "title",
    "project_name",
    "assignee_name",
    "department_name",
    "status_name",
    "due_date",
    "is_overdue",
    "closed_at",
    "closed_overdue",
    "days_overdue",
]
REPORT_COLUMNS_COMPACT = [
    "title",
    "project_name",
    "assignee_name",
    "department_name",
    "status_name",
    "due_date",
    "is_overdue",
    "closed_at",
]
REPORT_COLUMN_HEADERS: dict[str, str] = {
    "task_id": "ID задачи",
    "title": "Наименование задачи",
    "project_name": "Проект",
    "assignee_name": "Исполнитель",
    "department_name": "Департамент",
    "status_name": "Состояние",
    "due_date": "Срок",
    "is_overdue": "Просрочена",
    "closed_at": "Закрыта",
    "closed_overdue": "Закрыта с просрочкой",
    "days_overdue": "Дней просрочки",
}
REPORT_COLUMN_DEFAULT_WIDTHS: dict[str, int] = {
    "task_id": 10,
    "title": 40,
    "project_name": 28,
    "assignee_name": 28,
    "department_name": 28,
    "status_name": 20,
    "due_date": 20,
    "is_overdue": 14,
    "closed_at": 20,
    "closed_overdue": 22,
    "days_overdue": 16,
}


def utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _parse_datetime(value: str) -> datetime:
    return datetime.fromisoformat(value.replace("Z", "+00:00"))


def _calculate_days_overdue(
    *,
    due_date: datetime,
    closed_at: datetime | None,
    is_overdue: bool,
) -> int:
    if closed_at is not None and closed_at > due_date:
        return max(0, (closed_at.date() - due_date.date()).days)
    if is_overdue:
        return max(0, (utcnow().date() - due_date.date()).days)
    return 0


def build_task_report_rows(
    session: Session,
    *,
    date_from: date | None = None,
    date_to: date | None = None,
    project_id: int | None = None,
    department_id: int | None = None,
    assignee_id: int | None = None,
    workflow_status_id: int | None = None,
    overdue_only: bool = False,
    project_ids: set[int] | None = None,
    viewer_user_ids: set[int] | None = None,
) -> list[ReportTaskRow]:
    rows = report_repo.query_tasks_report(
        session,
        date_from=date_from,
        date_to=date_to,
        project_id=project_id,
        department_id=department_id,
        assignee_id=assignee_id,
        workflow_status_id=workflow_status_id,
        overdue_only=overdue_only,
        project_ids=project_ids,
        viewer_user_ids=viewer_user_ids,
    )

    result: list[ReportTaskRow] = []
    for row in rows:
        due_date = row[6]
        is_overdue = row[7]
        closed_at = row[8]
        closed_overdue = bool(closed_at and closed_at > due_date)
        days_overdue = _calculate_days_overdue(
            due_date=due_date,
            closed_at=closed_at,
            is_overdue=is_overdue,
        )
        result.append(
            ReportTaskRow(
                task_id=row[0],
                title=row[1],
                project_name=row[2],
                assignee_name=row[3],
                department_name=row[4],
                status_name=row[5],
                due_date=due_date.isoformat(),
                is_overdue=is_overdue,
                closed_at=closed_at.isoformat() if closed_at else None,
                closed_overdue=closed_overdue,
                days_overdue=days_overdue,
            )
        )
    return result


def build_discipline_rows(
    task_rows: list[ReportTaskRow],
) -> list[DisciplineRow]:
    rows: list[DisciplineRow] = []
    for row in task_rows:
        if not row.days_overdue:
            continue
        rows.append(
            DisciplineRow(
                department_name=row.department_name,
                project_name=row.project_name,
                assignee_name=row.assignee_name or "Не назначен",
                task_title=row.title,
                due_date=row.due_date,
                closed_at=row.closed_at,
                days_overdue=row.days_overdue,
            )
        )
    rows.sort(
        key=lambda item: (
            -item.days_overdue,
            item.department_name,
            item.project_name,
            item.assignee_name,
            item.task_title,
        )
    )
    return rows


def _normalize_export_template(template: str | None) -> str:
    if template == EXPORT_TEMPLATE_COMPACT:
        return EXPORT_TEMPLATE_COMPACT
    return EXPORT_TEMPLATE_FULL


def parse_report_columns(raw: str | None) -> list[str] | None:
    if not raw:
        return None
    values = [part.strip() for part in raw.split(",") if part.strip()]
    return values or None


def parse_report_column_widths(raw: str | None) -> list[int] | None:
    if not raw:
        return None
    widths: list[int] = []
    for chunk in raw.split(","):
        value = chunk.strip()
        if not value:
            continue
        try:
            widths.append(max(8, min(80, int(value))))
        except ValueError:
            continue
    return widths or None


def _resolve_report_columns(template: str, columns: list[str] | None) -> list[str]:
    if columns:
        resolved: list[str] = []
        for key in columns:
            if key in REPORT_COLUMN_HEADERS and key not in resolved:
                resolved.append(key)
        if resolved:
            return resolved
    if template == EXPORT_TEMPLATE_COMPACT:
        return REPORT_COLUMNS_COMPACT.copy()
    return REPORT_COLUMNS_FULL.copy()


def _report_headers(columns: list[str]) -> list[str]:
    return [REPORT_COLUMN_HEADERS[key] for key in columns]


def _report_row_values(row: ReportTaskRow, columns: list[str]) -> list[object]:
    value_map: dict[str, object] = {
        "task_id": row.task_id,
        "title": row.title,
        "project_name": row.project_name,
        "assignee_name": row.assignee_name,
        "department_name": row.department_name,
        "status_name": row.status_name,
        "due_date": row.due_date,
        "is_overdue": "Да" if row.is_overdue else "Нет",
        "closed_at": row.closed_at or "",
        "closed_overdue": "Да" if row.closed_overdue else "Нет",
        "days_overdue": row.days_overdue,
    }
    return [value_map[key] for key in columns]


def _resolve_report_column_widths(
    columns: list[str], widths: list[int] | None = None
) -> list[int]:
    defaults = [REPORT_COLUMN_DEFAULT_WIDTHS.get(key, 20) for key in columns]
    if not widths:
        return defaults
    merged = defaults.copy()
    for idx, width in enumerate(widths):
        if idx >= len(merged):
            break
        merged[idx] = max(8, min(80, width))
    return merged


def to_csv(
    rows: list[ReportTaskRow],
    *,
    template: str | None = None,
    columns: list[str] | None = None,
) -> str:
    template_key = _normalize_export_template(template)
    resolved_columns = _resolve_report_columns(template_key, columns)
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(_report_headers(resolved_columns))
    for row in rows:
        writer.writerow(_report_row_values(row, resolved_columns))
    return output.getvalue()


def _format_xlsx_header(ws, headers: list[str]) -> None:
    ws.append(headers)
    for col_idx, title in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=title)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def _set_xlsx_widths(ws, widths: list[int]) -> None:
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width


def _add_bar_sheet(
    wb: Workbook,
    *,
    title: str,
    headers: list[str],
    rows: list[list[object]],
    category_col: int = 1,
    value_col: int = 2,
) -> None:
    ws = wb.create_sheet(title=title)
    _format_xlsx_header(ws, headers)
    for row in rows:
        ws.append(row)
    for worksheet_row in ws.iter_rows(min_row=2):
        for cell in worksheet_row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    _set_xlsx_widths(ws, [36, 18, 18, 18])

    if not rows:
        return

    chart = BarChart()
    chart.title = title
    chart.y_axis.title = "Количество"
    chart.x_axis.title = headers[category_col - 1]
    values = Reference(
        ws,
        min_col=value_col,
        min_row=1,
        max_row=len(rows) + 1,
    )
    categories = Reference(
        ws,
        min_col=category_col,
        min_row=2,
        max_row=len(rows) + 1,
    )
    chart.add_data(values, titles_from_data=True)
    chart.set_categories(categories)
    chart.height = 9
    chart.width = 15
    ws.add_chart(chart, "F2")


def _add_pie_sheet(
    wb: Workbook,
    *,
    title: str,
    rows: list[list[object]],
) -> None:
    ws = wb.create_sheet(title=title)
    _format_xlsx_header(ws, ["Категория", "Количество"])
    for row in rows:
        ws.append(row)
    _set_xlsx_widths(ws, [36, 18])

    if not rows:
        return

    chart = PieChart()
    chart.title = title
    data = Reference(ws, min_col=2, min_row=1, max_row=len(rows) + 1)
    labels = Reference(ws, min_col=1, min_row=2, max_row=len(rows) + 1)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(labels)
    chart.height = 9
    chart.width = 15
    ws.add_chart(chart, "E2")


def _aggregate_totals(
    rows: list[ReportTaskRow],
    key_func: Callable[[ReportTaskRow], str],
) -> list[list[object]]:
    grouped: dict[str, tuple[int, int]] = {}
    for row in rows:
        key = key_func(row)
        total, overdue = grouped.get(key, (0, 0))
        grouped[key] = (total + 1, overdue + (1 if row.is_overdue else 0))
    return [
        [name, total, overdue]
        for name, (total, overdue) in sorted(grouped.items(), key=lambda item: item[1][0], reverse=True)
    ]


def _deadline_distribution(rows: list[ReportTaskRow]) -> list[list[object]]:
    today = utcnow().date()
    green = 0
    yellow = 0
    red = 0
    for row in rows:
        if row.is_overdue:
            red += 1
            continue
        due_date = _parse_datetime(row.due_date).date()
        delta = (due_date - today).days
        if delta <= 2:
            yellow += 1
        else:
            green += 1
    return [
        ["В срок", green],
        ["Критично", yellow],
        ["Просрочено", red],
    ]


def to_xlsx(
    rows: list[ReportTaskRow],
    *,
    template: str | None = None,
    columns: list[str] | None = None,
    column_widths: list[int] | None = None,
) -> bytes:
    template_key = _normalize_export_template(template)
    resolved_columns = _resolve_report_columns(template_key, columns)
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчёт по задачам"

    headers = _report_headers(resolved_columns)
    _format_xlsx_header(ws, headers)
    _set_xlsx_widths(
        ws,
        _resolve_report_column_widths(resolved_columns, column_widths),
    )

    for row in rows:
        ws.append(_report_row_values(row, resolved_columns))
    for worksheet_row in ws.iter_rows(min_row=2):
        for cell in worksheet_row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    by_department = _aggregate_totals(
        rows,
        key_func=lambda row: row.department_name or "Без департамента",
    )
    by_project = _aggregate_totals(
        rows,
        key_func=lambda row: row.project_name or "Без проекта",
    )
    by_assignee = _aggregate_totals(
        rows,
        key_func=lambda row: row.assignee_name or "Не назначен",
    )
    by_status = _aggregate_totals(
        rows,
        key_func=lambda row: row.status_name or "Без статуса",
    )
    _add_bar_sheet(
        wb,
        title="Диагр_Департаменты",
        headers=["Департамент", "Всего", "Просрочено"],
        rows=by_department,
        category_col=1,
        value_col=2,
    )
    _add_bar_sheet(
        wb,
        title="Диагр_Проекты",
        headers=["Проект", "Всего", "Просрочено"],
        rows=by_project,
        category_col=1,
        value_col=2,
    )
    _add_bar_sheet(
        wb,
        title="Диагр_Исполнители",
        headers=["Исполнитель", "Всего", "Просрочено"],
        rows=by_assignee,
        category_col=1,
        value_col=2,
    )
    _add_pie_sheet(
        wb,
        title="Диагр_Статусы",
        rows=[[row[0], row[1]] for row in by_status],
    )
    _add_pie_sheet(
        wb,
        title="Диагр_Сроки",
        rows=_deadline_distribution(rows),
    )

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def discipline_to_xlsx(rows: list[DisciplineRow]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Исполнительская дисциплина"
    headers = [
        "Департамент",
        "Проект",
        "Исполнитель",
        "Задача",
        "Срок",
        "Закрыта",
        "Дней просрочки",
    ]
    _format_xlsx_header(ws, headers)
    _set_xlsx_widths(ws, [28, 28, 28, 48, 20, 20, 16])

    for row in rows:
        ws.append(
            [
                row.department_name,
                row.project_name,
                row.assignee_name,
                row.task_title,
                row.due_date,
                row.closed_at or "",
                row.days_overdue,
            ]
        )
    for worksheet_row in ws.iter_rows(min_row=2):
        for cell in worksheet_row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def discipline_to_docx(rows: list[DisciplineRow]) -> bytes:
    def apply_tnr(run, *, size: int, bold: bool | None = None) -> None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("СПРАВКА ПО ИСПОЛНИТЕЛЬСКОЙ ДИСЦИПЛИНЕ")
    apply_tnr(title_run, size=14, bold=True)

    meta = document.add_paragraph(
        f"Дата формирования: {utcnow().strftime('%d.%m.%Y %H:%M UTC')}\n"
        f"Количество записей: {len(rows)}"
    )
    for run in meta.runs:
        apply_tnr(run, size=12)

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    column_widths = [Cm(3.1), Cm(3.0), Cm(3.1), Cm(6.8), Cm(2.4), Cm(2.4), Cm(2.2)]
    for idx, width in enumerate(column_widths):
        table.columns[idx].width = width

    headers = [
        "Департамент",
        "Проект",
        "Исполнитель",
        "Задача",
        "Срок",
        "Закрыта",
        "Дней просрочки",
    ]
    for idx, title in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = title
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if paragraph.runs:
            apply_tnr(paragraph.runs[0], size=11, bold=True)

    for row in rows:
        cells = table.add_row().cells
        values = [
            row.department_name,
            row.project_name,
            row.assignee_name,
            row.task_title,
            _parse_datetime(row.due_date).strftime("%d.%m.%Y"),
            _parse_datetime(row.closed_at).strftime("%d.%m.%Y") if row.closed_at else "",
            str(row.days_overdue),
        ]
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
                if idx in {4, 5, 6}
                else WD_PARAGRAPH_ALIGNMENT.LEFT
            )
            for run in paragraph.runs:
                apply_tnr(run, size=11)

    document.add_paragraph()
    signer = document.add_paragraph(
        "Ответственный: ____________________    Дата: ____________________"
    )
    signer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in signer.runs:
        apply_tnr(run, size=12)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
